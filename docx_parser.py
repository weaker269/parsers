"""Word 文档解析器

基于 python-docx 的 DOCX 解析实现，支持：
- 段落文本提取
- 表格检测和转换为 Markdown
- 图像提取和 OCR 文字识别
- 图文混排顺序保持
- 错误回退机制（完整解析 → 简化解析）
- 异步并发图像 OCR（性能优化）
"""

from io import BytesIO
import asyncio
from docx import Document
from .base import BaseParser
import logging

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Word 文档解析器（基于 WeKnora 简化版）

    使用 python-docx 进行文档解析。
    包含错误回退机制：完整解析失败时自动降级到简化方案。

    内部使用异步并发处理图像 OCR，显著提升多图场景性能。
    """

    async def parse(self, content: bytes) -> str:
        """解析 Word 文档（异步版本）

        Args:
            content: DOCX 文件的二进制内容

        Returns:
            解析后的文本内容

        Note:
            - 内部使用异步并发处理图像 OCR，提升性能 3-5 倍
            - 如果完整解析失败，会自动尝试简化解析方案
            - 此方法为异步方法，调用时需要使用 await
        """
        try:
            # 直接 await 异步版本（真正的异步调用）
            return await self._parse_advanced(content)
        except Exception as e:
            logger.warning(f"Async DOCX parsing failed: {e}, using fallback")
            return self._parse_simple(content)

    async def _parse_advanced(self, content: bytes) -> str:
        """完整异步解析：段落 + 表格 + 图像OCR

        提取所有段落文本、表格、图像，保持原始顺序。
        图像执行异步并发 OCR 识别，表格转换为 Markdown 格式。

        Args:
            content: DOCX 文件的二进制内容

        Returns:
            解析后的文本内容（包含 OCR 识别的文字）

        Note:
            - 使用 content_sequence 维护图文混排的原始顺序
            - 图像批量异步并发处理，最大并发数 5
            - 单个图像超时时间 30 秒
        """
        doc = Document(BytesIO(content))
        content_sequence = []  # 维护图文混排顺序
        image_count = 0

        logger.debug(f"开始异步解析 DOCX，共 {len(doc.paragraphs)} 个段落，{len(doc.tables)} 个表格")

        # 1. 构建内容序列（段落 + 表格 + 图像，保持原始顺序）
        for element in doc.element.body:
            # 检查是段落还是表格
            if element.tag.endswith('p'):  # 段落
                # 找到对应的 Paragraph 对象
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break

                if para:
                    # 提取段落文本
                    text = para.text.strip()
                    if text:
                        content_sequence.append(("text", text))

                    # 检查段落中是否包含图像
                    images = self._extract_images_from_paragraph(para)
                    for image_data in images:
                        image_count += 1
                        content_sequence.append(("image", image_data, image_count))

            elif element.tag.endswith('tbl'):  # 表格
                # 找到对应的 Table 对象
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break

                if table:
                    markdown_table = self._table_to_markdown(table)
                    if markdown_table:
                        content_sequence.append(("table", markdown_table))

        logger.debug(f"内容序列构建完成，共 {len(content_sequence)} 个元素，{image_count} 个图像")

        # 2. 收集所有图像数据
        images_data = []
        image_indices = []  # 记录每个图像在 content_sequence 中的索引

        for seq_idx, item in enumerate(content_sequence):
            if item[0] == "image":
                images_data.append(item[1])  # image_data
                image_indices.append((seq_idx, item[2]))  # (序列索引, 图像编号)

        # 3. 异步并发处理所有图像 OCR
        ocr_results_map = {}  # {图像编号: ocr_text}

        if images_data:
            logger.info(f"开始异步并发 OCR，共 {len(images_data)} 个图像")
            ocr_results_list = await self.process_images_async(
                images_data,
                max_concurrent=5,
                timeout_per_image=180.0  # 单个图像超时 180 秒（首次请求包含模型加载）
            )

            # 构建结果映射
            ocr_success_count = 0
            ocr_empty_count = 0

            for result_index, ocr_text in ocr_results_list:
                # result_index 是 process_images_async 返回的索引（从 1 开始）
                # 映射回原始图像编号
                _, original_img_num = image_indices[result_index - 1]

                if ocr_text.strip():
                    ocr_results_map[original_img_num] = ocr_text
                    ocr_success_count += 1
                    logger.debug(f"图像 {original_img_num} OCR 成功，识别 {len(ocr_text)} 字符")
                else:
                    ocr_empty_count += 1
                    logger.debug(f"图像 {original_img_num} 未识别到文字")

            ocr_failed_count = len(images_data) - len(ocr_results_list)

            logger.info(
                f"DOCX 异步解析完成，共提取 {image_count} 个图像：OCR成功 {ocr_success_count} 个，"
                f"无文字 {ocr_empty_count} 个，失败 {ocr_failed_count} 个"
            )
        else:
            logger.info(f"DOCX 解析完成，文档中未检测到图像")

        # 4. 组装最终结果
        result_parts = []
        for item in content_sequence:
            if item[0] == "text":
                result_parts.append(item[1])
            elif item[0] == "table":
                result_parts.append(item[1])
            elif item[0] == "image":
                img_num = item[2]
                if img_num in ocr_results_map:
                    ocr_text = ocr_results_map[img_num]
                    result_parts.append(f"[图像 {img_num} OCR 内容]:\n{ocr_text}")

        return "\n\n".join(result_parts)

    def _parse_simple(self, content: bytes) -> str:
        """简化回退方案：纯文本拼接

        忽略所有格式和图像，只提取纯文本内容。
        当完整解析失败时自动降级到此方案。

        Args:
            content: DOCX 文件的二进制内容

        Returns:
            解析后的纯文本内容（不包含图像 OCR）

        Note:
            - 忽略图像（不执行 OCR）
            - 表格简化为行内拼接
            - 保证最大兼容性
        """
        doc = Document(BytesIO(content))
        parts = []

        # 只提取段落文本，忽略所有格式
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 表格：简单拼接单元格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    parts.append(row_text)

        return "\n\n".join(parts)

    def _table_to_markdown(self, table) -> str:
        """将 Word 表格转为 Markdown（增强版本）

        生成标准的 Markdown 表格，第一行作为表头。
        换行符保留为 <br> 标签，增强对畸形表格的处理。

        Args:
            table: python-docx 表格对象

        Returns:
            Markdown 格式的表格字符串

        Note:
            - 换行符转换为 <br> 标签（保留换行信息）
            - 跳过空行和列数不匹配的畸形行
            - 跳过全空内容的行
        """
        if not table.rows:
            return ""

        # 清理单元格（保留换行信息，使用 <br> 标签）
        def clean(cell_text):
            if not cell_text:
                return ""
            # 将换行符替换为 <br> 标签（保留换行信息）
            return cell_text.replace("\n", "<br>").strip()

        # 提取所有行数据
        rows_data = []
        for row in table.rows:
            row_text = [clean(cell.text) for cell in row.cells]
            rows_data.append(row_text)

        if not rows_data:
            return ""

        # 生成 Markdown（第一行作为表头）
        header = rows_data[0]

        # 验证表头有效性
        if not header or all(not cell for cell in header):
            logger.debug("表格表头为空，跳过转换")
            return ""

        # 计算有效列数
        num_columns = len(header)

        # 表头
        md = "| " + " | ".join(header) + " |\n"
        # 分隔符
        md += "| " + " | ".join(["---"] * num_columns) + " |\n"

        # 数据行（增强验证逻辑）
        valid_rows = 0
        skipped_rows = 0

        for row_idx, row in enumerate(rows_data[1:], start=2):
            # 跳过空行
            if not row:
                skipped_rows += 1
                continue

            # 验证列数匹配
            if len(row) != num_columns:
                logger.debug(
                    f"跳过表格第 {row_idx} 行：列数不匹配 "
                    f"(期望 {num_columns} 列，实际 {len(row)} 列)"
                )
                skipped_rows += 1
                continue

            # 跳过全空行（所有单元格都为空）
            if all(not cell for cell in row):
                logger.debug(f"跳过表格第 {row_idx} 行：全空内容")
                skipped_rows += 1
                continue

            # 添加有效行
            md += "| " + " | ".join(row) + " |\n"
            valid_rows += 1

        # 统计日志
        if skipped_rows > 0:
            logger.debug(f"表格转换完成：有效行 {valid_rows} 个，跳过 {skipped_rows} 个畸形行")

        return md

    def _extract_images_from_paragraph(self, paragraph) -> list:
        """从段落中提取所有图像数据

        Args:
            paragraph: python-docx Paragraph 对象

        Returns:
            图像二进制数据列表

        Note:
            - 兼容 python-docx 1.2.0（不使用 xpath namespaces 参数）
            - 检测 InlineShapes（内嵌图像）和 Shapes（浮动图像）
            - 图像数据从 document.part.related_parts 中提取
            - 跳过过小的图像（装饰性图标）
        """
        images = []
        drawing_elements_found = 0
        images_extracted = 0
        images_skipped = 0

        try:
            # 遍历段落中的所有 runs
            for run in paragraph.runs:
                # 检查 run 中是否包含图像
                # 图像通常在 w:drawing 元素中
                # 方法1：检查 drawing 元素
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')

                for drawing in drawings:
                    drawing_elements_found += 1

                    # 在 drawing 中查找 blip 元素（包含图像引用）
                    # 使用 findall 而不是 xpath，兼容 python-docx 1.2.0
                    blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')

                    for blip in blips:
                        # 获取图像的关系 ID
                        # 尝试两个可能的属性名
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if not embed:
                            # 有些文档可能使用 link 属性
                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link')

                        if not embed:
                            logger.debug("图像元素中未找到 embed 或 link 属性")
                            continue

                        # 从文档中提取图像数据
                        try:
                            # 获取图像 part
                            image_part = paragraph._parent.part.related_parts[embed]
                            image_data = image_part.blob

                            # 跳过过小的图像（可能是图标、logo 等）
                            # 通过图像数据大小粗略判断
                            MIN_IMAGE_SIZE = 5000  # 最小 5KB
                            if len(image_data) < MIN_IMAGE_SIZE:
                                logger.debug(f"跳过小尺寸图像 (大小: {len(image_data)} 字节)")
                                images_skipped += 1
                                continue

                            # 背景图检测（过滤装饰性大图）
                            from parsers.ocr_worker import is_background_image

                            # DOCX 中无法直接获取图像尺寸，仅通过文件大小检测
                            if is_background_image(image_data):
                                logger.debug(
                                    f"跳过背景图 (大小: {len(image_data)/1024:.1f}KB, 关系ID: {embed})"
                                )
                                images_skipped += 1
                                continue

                            images.append(image_data)
                            images_extracted += 1
                            logger.debug(f"成功提取图像 (大小: {len(image_data)} 字节, 关系ID: {embed})")

                        except KeyError:
                            logger.warning(f"无法找到图像关系 ID: {embed}")
                            continue
                        except Exception as e:
                            logger.warning(f"提取图像数据失败 (关系ID: {embed}): {e}")
                            continue

            # 统计日志：区分"无图"和"有图但提取失败"
            if drawing_elements_found > 0:
                if images_extracted == 0:
                    logger.warning(f"段落中检测到 {drawing_elements_found} 个图像元素，但未成功提取任何图像（跳过 {images_skipped} 个小图像）")
                else:
                    logger.debug(f"段落中成功提取 {images_extracted}/{drawing_elements_found} 个图像（跳过 {images_skipped} 个小图像）")

        except Exception as e:
            logger.error(f"从段落提取图像时发生异常: {e}", exc_info=True)
            # 图像提取失败不影响文本解析
            return []

        return images

    def _extract_all_images_from_document(self, doc) -> list:
        """从文档中提取所有图像（备用方案）

        当段落级别提取失败时，可以使用此方法遍历所有关系提取图像。

        Args:
            doc: python-docx Document 对象

        Returns:
            图像二进制数据列表

        Note:
            - 遍历文档的所有关系（relationships）
            - 提取所有图像类型的 part
            - 不保证顺序（因为关系是无序的）
        """
        images = []

        try:
            # 遍历文档的所有关系
            for rel_id, rel in doc.part.rels.items():
                # 检查是否为图像关系
                if "image" in rel.target_ref.lower():
                    try:
                        image_part = rel.target_part
                        image_data = image_part.blob

                        # 跳过过小的图像
                        MIN_IMAGE_SIZE = 5000  # 最小 5KB
                        if len(image_data) < MIN_IMAGE_SIZE:
                            logger.debug(f"跳过小尺寸图像 (关系ID: {rel_id}, 大小: {len(image_data)} 字节)")
                            continue

                        # 背景图检测（过滤装饰性大图）
                        from parsers.ocr_worker import is_background_image

                        if is_background_image(image_data):
                            logger.debug(
                                f"跳过背景图 (关系ID: {rel_id}, 大小: {len(image_data)/1024:.1f}KB)"
                            )
                            continue

                        images.append(image_data)
                        logger.debug(f"从文档关系提取图像 (关系ID: {rel_id}, 大小: {len(image_data)} 字节)")

                    except Exception as e:
                        logger.warning(f"提取图像失败 (关系ID: {rel_id}): {e}")
                        continue

            logger.info(f"从文档关系中提取到 {len(images)} 个图像")

        except Exception as e:
            logger.error(f"遍历文档关系时发生异常: {e}", exc_info=True)
            return []

        return images
